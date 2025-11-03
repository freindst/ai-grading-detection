"""
Document Parser - Handle multiple file formats
Supports: PDF, DOCX, TXT, and images (with OCR)
"""

import os
from typing import Dict, Optional
from pathlib import Path


class DocumentParser:
    """Parse various document formats to extract text"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg', '.gif', '.webp']
    
    def parse_file(self, file_path: str) -> Dict:
        """
        Parse a file and extract text content
        
        Returns:
            Dict with success status, text content, and metadata
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return {
                    "success": False,
                    "error": "File not found",
                    "text": "",
                    "filename": str(file_path.name)
                }
            
            extension = file_path.suffix.lower()
            
            if extension not in self.supported_formats:
                return {
                    "success": False,
                    "error": f"Unsupported format: {extension}",
                    "text": "",
                    "filename": str(file_path.name)
                }
            
            # Route to appropriate parser
            if extension == '.pdf':
                text = self._parse_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                text = self._parse_docx(file_path)
            elif extension == '.txt':
                text = self._parse_txt(file_path)
            elif extension in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
                text = self._parse_image(file_path)
            else:
                return {
                    "success": False,
                    "error": f"Parser not implemented for {extension}",
                    "text": "",
                    "filename": str(file_path.name)
                }
            
            return {
                "success": True,
                "text": text,
                "filename": str(file_path.name),
                "format": extension,
                "size": file_path.stat().st_size
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "filename": str(file_path.name) if file_path else "unknown"
            }
    
    def _parse_pdf(self, file_path: Path) -> str:
        """Parse PDF file"""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(str(file_path))
            text = []
            
            for page in reader.pages:
                text.append(page.extract_text())
            
            return "\n".join(text)
        except ImportError:
            return "[Error: PyPDF2 not installed. Run: pip install PyPDF2]"
        except Exception as e:
            return f"[Error parsing PDF: {str(e)}]"
    
    def _parse_docx(self, file_path: Path) -> str:
        """Parse DOCX file"""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            text = []
            
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text.append(" | ".join(row_text))
            
            return "\n".join(text)
        except ImportError:
            return "[Error: python-docx not installed. Run: pip install python-docx]"
        except Exception as e:
            return f"[Error parsing DOCX: {str(e)}]"
    
    def _parse_txt(self, file_path: Path) -> str:
        """Parse plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # If all fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as f:
                return f.read().decode('utf-8', errors='ignore')
                
        except Exception as e:
            return f"[Error parsing TXT: {str(e)}]"
    
    def _parse_image(self, file_path: Path) -> str:
        """Parse image file using OCR"""
        try:
            from PIL import Image
            import pytesseract
            
            image = Image.open(str(file_path))
            text = pytesseract.image_to_string(image)
            
            if not text.strip():
                return "[Warning: No text detected in image]"
            
            return text
        except ImportError:
            return "[Error: PIL or pytesseract not installed. Run: pip install Pillow pytesseract]"
        except Exception as e:
            return f"[Error parsing image: {str(e)}]"
    
    def is_supported(self, filename: str) -> bool:
        """Check if file format is supported"""
        extension = Path(filename).suffix.lower()
        return extension in self.supported_formats

